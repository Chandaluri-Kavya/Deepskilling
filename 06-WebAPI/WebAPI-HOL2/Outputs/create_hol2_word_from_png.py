from docx import Document
from docx.shared import Inches
from pathlib import Path

output_dir = Path(__file__).resolve().parent
images_dir = output_dir / 'Documents'
doc_path = output_dir / 'WebAPI-HOL2-Screenshots.docx'

png_files = sorted(images_dir.glob('*.png'))

if not png_files:
    raise SystemExit('No PNG files found in Documents folder.')

print(f'Found {len(png_files)} PNG files.')


doc = Document()
doc.add_heading('WebAPI HOL2 Screenshot Outputs', level=1)

doc.add_paragraph('This document contains the exported PNG screenshots from the WebAPI HOL2 Outputs/Documents folder.')

doc.add_page_break()

for png in png_files:
    doc.add_heading(png.name, level=2)
    doc.add_paragraph(str(png))
    try:
        doc.add_picture(str(png), width=Inches(6))
    except Exception as exc:
        doc.add_paragraph(f'Failed to insert image: {exc}')
    doc.add_page_break()

print(f'Saving Word document to: {doc_path}')
doc.save(doc_path)
print('Done.')
